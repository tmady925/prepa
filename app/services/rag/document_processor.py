"""
Traitement des documents pour le RAG.
Supporte PDF (natif + scanné), Word, Images.
Chunking sémantique : respecte les frontières naturelles du contenu.
"""
import io
import re
import platform
import pytesseract
import fitz  # pymupdf
import cv2
import numpy as np
from PIL import Image
from docx import Document as DocxDocument

# Chemin Tesseract selon l'OS
if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Tesseract disponible ?
TESSERACT_AVAILABLE = True
try:
    pytesseract.get_tesseract_version()
except Exception:
    TESSERACT_AVAILABLE = False
    print("Tesseract non disponible — OCR désactivé")

# Taille des chunks sémantiques
CHUNK_MIN = 200
CHUNK_MAX = 1500
CHUNK_TARGET = 800


class DocumentProcessor:

    async def process(self, file_bytes: bytes, filename: str) -> dict:
        """
        Traite un document et retourne le texte extrait + les chunks sémantiques.
        """
        ext = filename.lower().split(".")[-1]

        try:
            if ext == "pdf":
                return await self._process_pdf(file_bytes)
            elif ext in ("docx", "doc"):
                return await self._process_docx(file_bytes)
            elif ext in ("png", "jpg", "jpeg", "tiff", "bmp"):
                return await self._process_image(file_bytes)
            else:
                return {"error": f"Format non supporté : {ext}"}
        except Exception as e:
            return {"error": str(e)}

    # ── PDF ───────────────────────────────────────────────────────────

    async def _process_pdf(self, file_bytes: bytes) -> dict:
        """Traite un PDF — natif ou scanné."""
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        has_ocr = False

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text").strip()

            if len(text) < 50:
                if TESSERACT_AVAILABLE:
                    has_ocr = True
                    text = self._ocr_pdf_page(page)
                else:
                    print(f"Page {page_num} scannée — Tesseract non disponible")

            if text:
                pages_text.append(text)

        doc.close()

        full_text = "\n\n".join(pages_text)
        full_text = self._clean_text(full_text)
        chunks = self._semantic_chunk(full_text)

        return {
            "text": full_text,
            "chunks": chunks,
            "page_count": len(pages_text),
            "has_ocr": has_ocr,
            "error": None,
        }

    # ── Word ──────────────────────────────────────────────────────────

    async def _process_docx(self, file_bytes: bytes) -> dict:
        """Traite un fichier Word."""
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        full_text = self._clean_text(full_text)
        chunks = self._semantic_chunk(full_text)

        return {
            "text": full_text,
            "chunks": chunks,
            "page_count": len(paragraphs) // 20 + 1,
            "has_ocr": False,
            "error": None,
        }

    # ── Image ─────────────────────────────────────────────────────────

    async def _process_image(self, file_bytes: bytes) -> dict:
        """Traite une image avec OCR."""
        if not TESSERACT_AVAILABLE:
            return {
                "text": "", "chunks": [], "page_count": 1,
                "has_ocr": False,
                "error": "Tesseract non disponible sur ce serveur",
            }

        image_array = np.frombuffer(file_bytes, np.uint8)
        img = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
        text = self._ocr_image(img)
        text = self._clean_text(text)
        chunks = self._semantic_chunk(text)

        return {
            "text": text, "chunks": chunks,
            "page_count": 1, "has_ocr": True, "error": None,
        }

    # ── Chunking sémantique ───────────────────────────────────────────

    def _semantic_chunk(self, text: str) -> list[str]:
        """
        Découpe le texte en respectant les frontières naturelles.

        Priorité des frontières (de la plus forte à la plus faible) :
        1. Titres de chapitres / exercices (CHAPITRE, EXERCICE, §...)
        2. Double saut de ligne (fin de paragraphe)
        3. Fin de phrase (. ! ?)
        4. Virgule ou point-virgule
        5. Espace (dernier recours)
        """
        if not text or len(text) < CHUNK_MIN:
            return [text] if text else []

        # 1. Détecte les frontières sémantiques fortes
        segments = self._split_by_semantic_boundaries(text)

        # 2. Fusionne les segments trop petits
        segments = self._merge_small_segments(segments)

        # 3. Coupe les segments trop grands
        chunks = []
        for seg in segments:
            if len(seg) <= CHUNK_MAX:
                chunks.append(seg)
            else:
                chunks.extend(self._split_large_segment(seg))

        # 4. Filtre les chunks vides ou trop courts
        chunks = [c.strip() for c in chunks if len(c.strip()) >= CHUNK_MIN]

        return chunks

    def _split_by_semantic_boundaries(self, text: str) -> list[str]:
        """
        Découpe aux frontières sémantiques fortes.
        Détecte : titres, numéros d'exercices, sections, sous-titres.
        """
        # Patterns de frontières fortes
        boundary_patterns = [
            # Titres de chapitres
            r'\n(?=CHAPITRE\s+\w)',
            r'\n(?=Chapitre\s+\w)',
            r'\n(?=PARTIE\s+\w)',
            r'\n(?=Partie\s+\w)',

            # Exercices et problèmes
            r'\n(?=Exercice\s*\d)',
            r'\n(?=EXERCICE\s*\d)',
            r'\n(?=Problème\s*\d)',
            r'\n(?=PROBLEME\s*\d)',
            r'\n(?=Activité\s*\d)',

            # Questions numérotées
            r'\n(?=\d+[°\.\)]\s+[A-ZÀÂÉÈÊÎÔÙÛ])',

            # Sections avec tirets ou points
            r'\n(?=[IVX]+[\.:\-]\s+[A-ZÀÂÉÈÊÎÔÙÛ])',
            r'\n(?=[A-Z][\.:\-]\s+[A-ZÀÂÉÈÊÎÔÙÛ])',

            # Définitions, théorèmes, propriétés
            r'\n(?=Définition\s*[:\.]\s)',
            r'\n(?=DEFINITION\s*[:\.]\s)',
            r'\n(?=Théorème\s*[:\.]\s)',
            r'\n(?=Propriété\s*[:\.]\s)',
            r'\n(?=Remarque\s*[:\.]\s)',
            r'\n(?=Méthode\s*[:\.]\s)',
            r'\n(?=Corrigé\s*[:\.]\s)',
            r'\n(?=CORRIGE\s*[:\.]\s)',
            r'\n(?=Solution\s*[:\.]\s)',
        ]

        # Combine tous les patterns
        combined = '|'.join(boundary_patterns)

        # Découpe
        parts = re.split(combined, text)
        return [p.strip() for p in parts if p.strip()]

    def _merge_small_segments(self, segments: list[str]) -> list[str]:
        """Fusionne les segments trop petits avec le suivant."""
        if not segments:
            return []

        merged = []
        current = segments[0]

        for seg in segments[1:]:
            if len(current) < CHUNK_MIN:
                # Fusionne avec le suivant
                current = current + "\n\n" + seg
            elif len(current) + len(seg) <= CHUNK_TARGET:
                # Fusionne si ça reste dans la taille cible
                current = current + "\n\n" + seg
            else:
                merged.append(current)
                current = seg

        if current:
            merged.append(current)

        return merged

    def _split_large_segment(self, text: str) -> list[str]:
        """
        Coupe un segment trop grand aux frontières faibles.
        Respecte les phrases complètes.
        """
        chunks = []
        current = ""

        # Découpe par phrases
        sentences = re.split(r'(?<=[.!?])\s+', text)

        for sentence in sentences:
            if len(current) + len(sentence) <= CHUNK_MAX:
                current = current + " " + sentence if current else sentence
            else:
                if current:
                    chunks.append(current.strip())
                # Si la phrase seule dépasse CHUNK_MAX → coupe par mots
                if len(sentence) > CHUNK_MAX:
                    words = sentence.split()
                    current = ""
                    for word in words:
                        if len(current) + len(word) + 1 <= CHUNK_MAX:
                            current = current + " " + word if current else word
                        else:
                            if current:
                                chunks.append(current.strip())
                            current = word
                else:
                    current = sentence

        if current:
            chunks.append(current.strip())

        return chunks

    # ── OCR ───────────────────────────────────────────────────────────

    def _ocr_pdf_page(self, page) -> str:
        """OCR sur une page PDF scannée."""
        mat = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        image_array = np.frombuffer(img_bytes, np.uint8)
        img = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
        return self._ocr_image(img)

    def _ocr_image(self, img: np.ndarray) -> str:
        """OCR avec prétraitement pour améliorer les scans flous."""
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        denoised = cv2.fastNlMeansDenoising(gray, h=10)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(denoised)
        binary = cv2.adaptiveThreshold(
            enhanced, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )
        pil_img = Image.fromarray(binary)
        text = pytesseract.image_to_string(
            pil_img,
            lang="fra",
            config="--psm 6 --oem 3"
        )
        return text.strip()

    # ── Nettoyage ─────────────────────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        """Nettoie le texte extrait."""
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()


document_processor = DocumentProcessor()